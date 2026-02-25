"""
Core RAG engine — retrieval, reranking, LangGraph agent, and semantic cache.
Extracted from the notebook for production use.
"""
import os
import json
import time
import hashlib
from typing import TypedDict, List, Dict, Any, Optional
from collections import defaultdict

import numpy as np
from langchain_openai import AzureChatOpenAI, AzureOpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langgraph.graph import StateGraph, END, START
from rank_bm25 import BM25Okapi

import config


# ══════════════════════════════════════════════════════════════
#  Globals — initialised by init_engine()
# ══════════════════════════════════════════════════════════════
llm: AzureChatOpenAI = None
embeddings: AzureOpenAIEmbeddings = None
vector_store: FAISS = None
bm25_index: BM25Okapi = None
all_chunks: List[Document] = []
agent_graph = None
query_cache = None


# ══════════════════════════════════════════════════════════════
#  Semantic Query Cache
# ══════════════════════════════════════════════════════════════
class SemanticQueryCache:
    def __init__(self, embedding_fn, threshold=0.92, max_size=50):
        self.embedding_fn = embedding_fn
        self.threshold = threshold
        self.max_size = max_size
        self.cache: list = []
        self.stats = {"hits": 0, "misses": 0}

    def _cosine(self, a, b):
        a, b = np.array(a), np.array(b)
        return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-10))

    def get(self, query: str, method: str = "default"):
        emb = self.embedding_fn(query)
        best_score, best = 0.0, None
        for c_emb, c_q, c_m, c_r in self.cache:
            if c_m != method:
                continue
            sim = self._cosine(emb, c_emb)
            if sim > best_score:
                best_score, best = sim, (c_q, c_r)
        if best_score >= self.threshold and best:
            self.stats["hits"] += 1
            return {"result": best[1], "matched_query": best[0], "similarity": best_score}
        self.stats["misses"] += 1
        return None

    def put(self, query: str, method: str, result: dict):
        emb = self.embedding_fn(query)
        if len(self.cache) >= self.max_size:
            self.cache.pop(0)
        self.cache.append((emb, query, method, result))


# ══════════════════════════════════════════════════════════════
#  Document loading & chunking
# ══════════════════════════════════════════════════════════════
# ── CORE CAPABILITY 1: Ingest and Process Company Documents ──
def load_and_chunk() -> List[Document]:
    """Load documents from disk and split into chunks."""
    os.makedirs(config.DOC_DIR, exist_ok=True)
    documents = []
    for fname in sorted(os.listdir(config.DOC_DIR)):
        fpath = os.path.join(config.DOC_DIR, fname)
        if os.path.isfile(fpath):
            with open(fpath, "r", encoding="utf-8") as f:
                content = f.read()
            documents.append(Document(page_content=content, metadata={"source": fname}))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = []
    for doc in documents:
        doc_chunks = splitter.split_documents([doc])
        for i, chunk in enumerate(doc_chunks):
            chunk.metadata["chunk_id"] = i
        chunks.extend(doc_chunks)
    return chunks


# ══════════════════════════════════════════════════════════════
#  CORE CAPABILITY 2: Retrieve Relevant Information for a Query
# ══════════════════════════════════════════════════════════════
def retrieve_dense(query: str, k: int = 5) -> List[Document]:
    return vector_store.similarity_search(query, k=k)


def retrieve_bm25(query: str, k: int = 5) -> List[Document]:
    scores = bm25_index.get_scores(query.lower().split())
    top_idx = np.argsort(scores)[::-1][:k]
    return [all_chunks[i] for i in top_idx]


def reciprocal_rank_fusion(dense_docs, sparse_docs, k=60):
    doc_scores, doc_map = defaultdict(float), {}
    for rank, doc in enumerate(dense_docs):
        key = hashlib.md5(doc.page_content[:200].encode()).hexdigest()
        doc_scores[key] += 1.0 / (rank + k)
        doc_map[key] = doc
    for rank, doc in enumerate(sparse_docs):
        key = hashlib.md5(doc.page_content[:200].encode()).hexdigest()
        doc_scores[key] += 1.0 / (rank + k)
        doc_map[key] = doc
    sorted_keys = sorted(doc_scores, key=lambda x: doc_scores[x], reverse=True)
    return [doc_map[k] for k in sorted_keys]


def retrieve_hybrid(query: str, k: int = 5) -> List[Document]:
    dense = retrieve_dense(query, k=k * 2)
    sparse = retrieve_bm25(query, k=k * 2)
    return reciprocal_rank_fusion(dense, sparse)[:k]


def llm_rerank(query: str, documents: List[Document], top_n: int = 5) -> List[Document]:
    if not documents:
        return []
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a relevance judge. Rate relevance 0-10. Reply with ONLY a number."),
        ("human", "Query: {query}\n\nDocument:\n{document}\n\nRelevance score (0-10):"),
    ])
    scored = []
    for doc in documents[:10]:
        try:
            resp = llm.invoke(prompt.format_messages(query=query, document=doc.page_content[:500]))
            score = float(resp.content.strip())
        except Exception:
            score = 5.0
        scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for _, d in scored[:top_n]]


def retrieve_hybrid_reranked(query: str, k: int = 5) -> List[Document]:
    return llm_rerank(query, retrieve_hybrid(query, k=k * 2), top_n=k)


# ══════════════════════════════════════════════════════════════
#  CORE CAPABILITY 3: Generate a Concise, Factual Answer
#  Grounded in the Documents
# ══════════════════════════════════════════════════════════════
def generate_answer_simple(query: str, docs: List[Document]) -> str:
    context = "\n\n---\n\n".join(
        [f"[Source: {d.metadata.get('source','unknown')}]\n{d.page_content}" for d in docs]
    )
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are an AI Leadership Insight Agent for ACME Corporation. "
                   "Answer ONLY from the provided context. Cite sources using [Source: filename]. "
                   "Be concise, factual, and structured."),
        ("human", "Context:\n{context}\n\nQuestion: {question}"),
    ])
    return llm.invoke(prompt.format_messages(context=context, question=query)).content


# ══════════════════════════════════════════════════════════════
#  LangGraph Agent (Method 5)
# ══════════════════════════════════════════════════════════════
class AgentState(TypedDict):
    question: str
    question_type: str
    sub_queries: List[str]
    retrieved_docs: List[Document]
    reranked_docs: List[Document]
    answer: str
    sources: List[str]
    confidence: float
    quality_pass: bool
    iteration: int


def _node_analyze_and_decompose(state: AgentState) -> dict:
    """Merged node: classify question AND decompose in a single LLM call.
    Saves one full LLM round-trip vs two separate calls."""
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a query analyzer for a corporate leadership AI assistant.
1. Classify the question as FACTUAL, ANALYTICAL, or STRATEGIC.
2. If FACTUAL, return the original question as the only sub-query.
   If ANALYTICAL or STRATEGIC, decompose into 2-3 focused sub-queries.

Reply with ONLY a JSON object: {{"type": "FACTUAL", "sub_queries": ["..."]}}    """),
        ("human", "{question}"),
    ])
    resp = llm.invoke(prompt.format_messages(question=state["question"]))
    try:
        r = json.loads(resp.content.strip())
        q_type = r.get("type", "ANALYTICAL").upper()
        subs = r.get("sub_queries", [state["question"]])
        if q_type not in ("FACTUAL", "ANALYTICAL", "STRATEGIC"):
            q_type = "ANALYTICAL"
        if not isinstance(subs, list) or not subs:
            subs = [state["question"]]
    except Exception:
        q_type = "ANALYTICAL"
        subs = [state["question"]]
    return {"question_type": q_type, "sub_queries": subs}


def _node_retrieve(state: AgentState) -> dict:
    all_docs, seen = [], set()
    for sq in state["sub_queries"]:
        for d in retrieve_hybrid(sq, k=5):
            key = hashlib.md5(d.page_content[:200].encode()).hexdigest()
            if key not in seen:
                seen.add(key)
                all_docs.append(d)
    return {"retrieved_docs": all_docs}


def _node_batch_rerank(state: AgentState) -> dict:
    """Batch rerank: score all documents in a single LLM call.
    Reduces ~10 sequential LLM calls to just 1."""
    docs = state["retrieved_docs"]
    if not docs:
        return {"reranked_docs": []}
    if len(docs) <= 5:
        return {"reranked_docs": docs}

    doc_list = "\n\n".join(
        f"[{i}] {d.page_content[:400]}" for i, d in enumerate(docs)
    )
    prompt = ChatPromptTemplate.from_messages([
        ("system", "You are a relevance judge. Given a query and numbered document chunks, "
                   "return the indices of the 5 most relevant documents ordered by relevance (most relevant first). "
                   "Reply with ONLY a JSON list of integers, e.g. [3, 0, 7, 1, 4]"),
        ("human", "Query: {query}\n\nDocuments:\n{documents}\n\nTop 5 most relevant indices:"),
    ])
    resp = llm.invoke(prompt.format_messages(query=state["question"], documents=doc_list))
    try:
        indices = json.loads(resp.content.strip())
        if not isinstance(indices, list):
            indices = list(range(min(5, len(docs))))
        valid = [i for i in indices if isinstance(i, int) and 0 <= i < len(docs)]
        if not valid:
            valid = list(range(min(5, len(docs))))
        return {"reranked_docs": [docs[i] for i in valid[:5]]}
    except Exception:
        return {"reranked_docs": docs[:5]}


def _node_synthesize(state: AgentState) -> dict:
    context = "\n\n---\n\n".join(
        [f"[Source: {d.metadata.get('source','unknown')}]\n{d.page_content}" for d in state["reranked_docs"]]
    )
    system = (
        "You are an AI Leadership Insight Agent for ACME Corporation.\n"
        "RULES: Answer ONLY from context. Cite [Source: filename]. Use headers and bullets.\n"
        "FORMAT: ## Key Findings\\n<answer>\\n## Supporting Details\\n<details>\\n## Sources\\n<sources>"
    )
    prompt = ChatPromptTemplate.from_messages([("system", system), ("human", "Context:\n{context}\n\nQuestion: {question}")])
    resp = llm.invoke(prompt.format_messages(context=context, question=state["question"]))
    sources = list(set(d.metadata.get("source", "unknown") for d in state["reranked_docs"]))
    return {"answer": resp.content, "sources": sources}


def _node_quality(state: AgentState) -> dict:
    prompt = ChatPromptTemplate.from_messages([
        ("system", 'Check answer quality. Reply JSON: {{"pass": true/false, "confidence": 0.0-1.0, "reason": "..."}}'),
        ("human", "Question: {question}\nAnswer: {answer}\nSources: {sources}"),
    ])
    resp = llm.invoke(prompt.format_messages(
        question=state["question"], answer=state["answer"], sources=", ".join(state["sources"])
    ))
    try:
        r = json.loads(resp.content.strip())
        return {"quality_pass": r.get("pass", True), "confidence": float(r.get("confidence", 0.7)),
                "iteration": state.get("iteration", 0) + 1}
    except Exception:
        return {"quality_pass": True, "confidence": 0.7, "iteration": state.get("iteration", 0) + 1}


def _should_retry(state: AgentState) -> str:
    if not state.get("quality_pass", True) and state.get("iteration", 0) < 2:
        return "retry"
    return "finish"


def _build_agent():
    wf = StateGraph(AgentState)
    wf.add_node("analyze_and_decompose", _node_analyze_and_decompose)
    wf.add_node("retrieve_documents", _node_retrieve)
    wf.add_node("batch_rerank", _node_batch_rerank)
    wf.add_node("synthesize_answer", _node_synthesize)
    wf.add_node("quality_check", _node_quality)
    wf.add_edge(START, "analyze_and_decompose")
    wf.add_edge("analyze_and_decompose", "retrieve_documents")
    wf.add_edge("retrieve_documents", "batch_rerank")
    wf.add_edge("batch_rerank", "synthesize_answer")
    wf.add_edge("synthesize_answer", "quality_check")
    wf.add_conditional_edges("quality_check", _should_retry, {"retry": "synthesize_answer", "finish": END})
    return wf.compile()


# ══════════════════════════════════════════════════════════════
#  Public API — run a query through any method
# ══════════════════════════════════════════════════════════════
METHODS = {
    "naive_dense": retrieve_dense,
    "bm25": retrieve_bm25,
    "hybrid": retrieve_hybrid,
    "hybrid_reranked": retrieve_hybrid_reranked,
}


def ask(question: str, method: str = "agentic") -> dict:
    """Run a question through the specified RAG method and return the result."""
    # 1. Check cache
    cached = query_cache.get(question, method=method)
    if cached:
        return {**cached["result"], "cache_hit": True, "matched_query": cached["matched_query"],
                "similarity": cached["similarity"]}

    t0 = time.time()

    if method == "agentic":
        state = {
            "question": question, "question_type": "", "sub_queries": [],
            "retrieved_docs": [], "reranked_docs": [], "answer": "",
            "sources": [], "confidence": 0.0, "quality_pass": False, "iteration": 0,
        }
        res = agent_graph.invoke(state)
        result = {
            "answer": res["answer"], "sources": res["sources"],
            "confidence": res.get("confidence", 0.0),
            "question_type": res.get("question_type", ""),
            "method": "agentic",
        }
    else:
        retriever = METHODS.get(method, retrieve_hybrid)
        docs = retriever(question, k=config.TOP_K)
        answer = generate_answer_simple(question, docs)
        result = {
            "answer": answer,
            "sources": list(set(d.metadata["source"] for d in docs)),
            "confidence": 0.8,
            "question_type": "N/A",
            "method": method,
        }

    result["latency_seconds"] = round(time.time() - t0, 2)
    result["cache_hit"] = False
    result["matched_query"] = None

    # 2. Store in cache
    query_cache.put(question, method=method, result=result)
    return result


# ══════════════════════════════════════════════════════════════
#  Initialisation — call once at startup
# ══════════════════════════════════════════════════════════════
def init_engine():
    """Initialise LLM, embeddings, vector store, BM25, agent graph, and cache."""
    global llm, embeddings, vector_store, bm25_index, all_chunks, agent_graph, query_cache

    print("[RAG] Initialising models…")
    llm = AzureChatOpenAI(
        azure_deployment=config.AZURE_CHAT_MODEL,
        azure_endpoint=config.AZURE_OPENAI_ENDPOINT,
        api_key=config.AZURE_OPENAI_API_KEY,
        api_version=config.AZURE_OPENAI_API_VERSION,
        max_tokens=2048,
    )
    embeddings = AzureOpenAIEmbeddings(
        model=config.AZURE_EMBEDDING_MODEL,
        azure_endpoint=config.AZURE_OPENAI_ENDPOINT,
        api_key=config.AZURE_OPENAI_API_KEY,
        api_version=config.AZURE_OPENAI_API_VERSION,
    )

    print("[RAG] Loading documents…")
    all_chunks = load_and_chunk()
    print(f"[RAG] {len(all_chunks)} chunks created")

    print("[RAG] Building FAISS index…")
    vector_store = FAISS.from_documents(all_chunks, embeddings)

    print("[RAG] Building BM25 index…")
    tokenized = [c.page_content.lower().split() for c in all_chunks]
    bm25_index = BM25Okapi(tokenized)

    print("[RAG] Compiling LangGraph agent…")
    agent_graph = _build_agent()

    print("[RAG] Initialising semantic cache…")
    query_cache = SemanticQueryCache(
        embedding_fn=embeddings.embed_query,
        threshold=config.CACHE_SIMILARITY_THRESHOLD,
        max_size=config.CACHE_MAX_SIZE,
    )

    print("[RAG] Engine ready ✓")


def get_doc_info() -> list:
    """Return info about loaded documents."""
    from collections import Counter
    counts = Counter(c.metadata["source"] for c in all_chunks)
    result = []
    for fname in sorted(counts):
        fpath = os.path.join(config.DOC_DIR, fname)
        size = os.path.getsize(fpath) if os.path.exists(fpath) else 0
        result.append({"filename": fname, "size_bytes": size, "chunks": counts[fname]})
    return result


def add_document(filename: str, content: bytes, ext: str) -> dict:
    """
    Process an uploaded PDF or DOCX file, add to knowledge base, and rebuild indexes.
    Uses pymupdf4llm for PDF and python-docx for Word.
    """
    global vector_store, bm25_index, all_chunks

    # Save file to disk
    fpath = os.path.join(config.DOC_DIR, filename)
    with open(fpath, "wb") as f:
        f.write(content)

    # Extract text
    if ext == ".pdf":
        import pymupdf4llm
        text = pymupdf4llm.to_markdown(fpath)
    elif ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(fpath)
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not text.strip():
        os.remove(fpath)
        raise ValueError("No text could be extracted from the file.")

    # Chunk the new document
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    new_doc = Document(page_content=text, metadata={"source": filename})
    new_chunks = splitter.split_documents([new_doc])
    for i, chunk in enumerate(new_chunks):
        chunk.metadata["chunk_id"] = i

    # Add to global chunks
    all_chunks.extend(new_chunks)

    # Rebuild FAISS index (full rebuild — simple and safe)
    print(f"[RAG] Rebuilding indexes with {len(all_chunks)} chunks (+{len(new_chunks)} from {filename})…")
    vector_store = FAISS.from_documents(all_chunks, embeddings)

    # Rebuild BM25 index
    tokenized = [c.page_content.lower().split() for c in all_chunks]
    bm25_index = BM25Okapi(tokenized)

    # Clear cache (new knowledge available)
    query_cache.cache.clear()
    query_cache.stats = {"hits": 0, "misses": 0}

    print(f"[RAG] Document '{filename}' added — {len(new_chunks)} chunks. Total: {len(all_chunks)} chunks.")
    return {
        "status": "success",
        "filename": filename,
        "chunks_added": len(new_chunks),
        "total_chunks": len(all_chunks),
    }
